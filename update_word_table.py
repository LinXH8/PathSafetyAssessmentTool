# -*- coding: utf-8 -*-
"""Update Table 1 (section 2.3) in the Word user guide to match the markdown."""

import docx
from docx.oxml.ns import qn
import copy

DOC_PATH = r"c:\Users\23010975\Documents\GitHub\PathSafetyAssessmentTool\USER_GUIDE_(UPDATED)V3.docx"

NEW_ROWS = [
    ("1",  "Area Type",                               "Suburban",               "GIS fallback using LTA parking zone and URA land use layers; default when no urban/industrial/recreational polygon match"),
    ("2",  "Facility Type",                            "Sidewalk",               "Most common facility type; overridden by CV image analysis"),
    ("3",  "Facility Access",                          "Adequate",               "Most common; code Inadequate only when congestion or blockages force users onto alternative routes"),
    ("4",  "Loose or Slippery Surface",                "Not Present",            "Most common; GIS-assisted via PATH defects layer (algae); CV detection not yet ready"),
    ("5",  "Tram or Train Rails",                      "Not Present",            "No tram or train rail paths/roads exist in Singapore"),
    ("6",  "Major Surface Deformation or Drain Opening","Not Present",           "Most common; GIS-assisted via PATH defects layer (tile cracks, tactile tile cracks, uneven pavement)"),
    ("7",  "Fixed Obstacle on Facility",               "Not Present",            "Most common; overridden by CV image analysis"),
    ("8",  "Non-Fixed Obstacle on Facility",           "Not Present",            "Most common; overridden by CV image analysis"),
    ("9",  "Line of Sight",                            "Not coded",              "Not yet in scoring specification; requires manual assessment of visibility along the path"),
    ("10", "Delineation",                              "Not Present",            "Default not present; CV codes first, then overwritten if PATH provides faded marking data"),
    ("11", "Light Segregation",                        "Present",                "All off-road paths are expected to have kerb segregation; confirmed by CV and logic rules"),
    ("12", "Facility Width per Direction",             "Narrow",                 "GIS-derived from LiDAR path width data; default when no width data available"),
    ("13", "Flow Direction",                           "One Way",                "Most on-road cycling lanes are one-way; all off-road paths default to two-way"),
    ("14", "Width Restriction",                        "Not Present",            "Most common; logic rule sets Present when a fixed or non-fixed obstacle is detected"),
    ("15", "Adjacent Road Lane 0–1m",             "See image logic",        "CV-derived; present when path edge is ≤1m from road kerb with no protective barrier"),
    ("16", "Adjacent Vehicle Parking 0–1m",       "Not Present",            "Most common; GIS-derived from URA car park lot layer and GDM lane markings"),
    ("17", "Adjacent Severe Hazard 0–1m",         "Not Present",            "Rare in Singapore; requires manual assessment when waterways or drops >60cm are present with no barrier"),
    ("18", "Adjacent Object/Level Change 0–1m",   "Mirrors Adj. Road 0–1m", "Co-occurs with adjacent road lane; inferred by CV and logic rules"),
    ("19", "Adjacent Sidewalk 0–1m",              "Present",                "Assumed present in urban context; derived by CV and logic rules"),
    ("20", "Adjacent Road Lane 1–3m",             "See image logic",        "CV-derived; present when road is 1–3m from path with no protective barrier"),
    ("21", "Adjacent Vehicle Parking 1–3m",       "Not Present",            "Most common; GIS-derived from URA car park lot layer"),
    ("22", "Adjacent Severe Hazard 1–3m",         "Not Present",            "Rare in Singapore; requires manual assessment when severe hazards are 1–3m from path"),
    ("23", "Adjacent Object/Level Change 1–3m",   "Mirrors Adj. Road 1–3m", "Co-occurs with adjacent road lane; inferred by CV and logic rules"),
    ("24", "Adjacent Sidewalk 1–3m",              "Not Present",            "Not present by default; CV and logic rules detect when split paths with >1m green verge are present"),
    ("25", "Grade",                                    "< 5 Degrees",            "Most paths in Singapore are flat; GIS-derived from LiDAR 3D point cloud data"),
    ("26", "Curvature",                                "No Sharp Turn",          "Most common; GIS-derived using path geometry and curvature radius calculation"),
    ("27", "Street Lighting",                          "Present",                "Most urban paths in Singapore have adequate street lighting"),
    ("28", "Pedestrian Crossing",                      "Not Present",            "Most common; GIS-derived from MRT exit and bus stop layers; manual coding where needed"),
    ("29", "Intersecting Bicycle Facility",            "Not Present",            "Most common; GIS-derived using pedestrian crossing layer within 5m radius"),
    ("30", "Intersection Approach",                    "Separate/NA",            "Most off-road cycling paths remain separate from traffic at intersections"),
    ("31", "Intersection or Road Crossing",            "Not Present",            "Most common; overridden by CV when crossings are detected"),
    ("32", "Crossing Facility",                        "Not Present",            "Most common; overridden by CV when zebra crossings or bicycle crossings are detected"),
    ("33", "Number of Lanes – Adjacent Road",    "1 per Direction/NA",     "Most common; GIS-derived from LiDAR kerbline layer"),
    ("34", "Number of Lanes – Intersecting Road","1 per Direction/NA",     "Default when no intersection is present; GIS-derived from LiDAR kerbline layer when intersection exists"),
    ("35", "Property Access",                          "Not Present",            "Most common; CV and GIS auto-coding pending"),
    ("36", "Peak Pedestrian Flow",                     "Low",                    "GIS-derived from MRT exit, bus stop, and AMG sensor count layers; default when no count data available"),
    ("37", "Peak Bicycle/LV Traffic Flow",             "Low",                    "GIS-derived from AMG 24/7 sensor count data; default when no count data available"),
    ("38", "Observed Proportion of Cargo Bikes",       "Low",                    "Most common; ePMD and PAB in Singapore are under 20kg, classifying them as low cargo proportion"),
    ("39", "Bicycle/LV Speed – Average",         "< 20 km/h",             "Most common cycling speed on Singapore shared paths and CPNs"),
    ("40", "Bicycle/LV Speed Differential",            "< 10 km/h",             "Most common; low speed variation typical on Singapore shared paths"),
    ("41", "Road AADT",                                "—",                 "Must be coded manually; no GIS auto-code path; ERP2 data integration pending"),
    ("42", "Heavy Vehicle Flow",                       "Low",                    "Most paths are away from heavy traffic; GIS-derived from bus lane marking layer (GDM)"),
    ("43", "Road Speed Limit",                         "—",                 "GIS-derived from GDM speed limit layer; used for analysis only and does not affect risk scores"),
    ("44", "Road Operating Speed (mean)",              "—",                 "GIS-derived from LinkID layer and ERP2 speed data; required when adjacent road or road crossing is present"),
]


def set_cell_text(cell, text):
    """Replace all runs in the first paragraph of a cell with a single run."""
    para = cell.paragraphs[0]
    # Preserve the paragraph formatting but clear all runs
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def main():
    doc = docx.Document(DOC_PATH)
    table = doc.tables[1]

    assert len(table.rows) == 45, f"Expected 45 rows (header + 44 data), got {len(table.rows)}"

    for i, (num, attr, default, reason) in enumerate(NEW_ROWS):
        row = table.rows[i + 1]  # skip header row
        cells = row.cells
        set_cell_text(cells[0], num)
        set_cell_text(cells[1], attr)
        set_cell_text(cells[2], default)
        set_cell_text(cells[3], reason)

    doc.save(DOC_PATH)
    print("Done. Saved to", DOC_PATH)


if __name__ == "__main__":
    main()
