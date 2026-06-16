"""
Generates DEVELOPER_GUIDE_(UPDATED)V3.docx from the ten developer-guide
markdown files served by the PSAT Help page (Documentation & Guides > Developer Guide tab).
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PUBLIC_DIR = Path(__file__).parent.parent / "frontend" / "public"
OUTPUT     = Path(__file__).parent.parent / "DEVELOPER_GUIDE_(UPDATED)V3.docx"

# Paths match the `path` entries in DeveloperGuide.tsx DOCS_LIST
MD_FILES = [
    PUBLIC_DIR / "README.md",
    PUBLIC_DIR / "docs" / "developer" / "installation.md",
    PUBLIC_DIR / "docs" / "developer" / "architecture.md",
    PUBLIC_DIR / "docs" / "developer" / "api-reference.md",
    PUBLIC_DIR / "docs" / "developer" / "cv-pipeline.md",
    PUBLIC_DIR / "docs" / "developer" / "scoring.md",
    PUBLIC_DIR / "docs" / "developer" / "frontend.md",
    PUBLIC_DIR / "docs" / "developer" / "common-issues.md",
    PUBLIC_DIR / "docs" / "developer" / "contributing.md",
    PUBLIC_DIR / "docs" / "developer" / "dev-jira.md",
]

# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(0)


def apply_inline_formatting(run_parent, raw: str):
    """Parse inline **bold**, `code`, and plain text; render links as plain text."""
    raw = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", raw)
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(raw)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = run_parent.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = run_parent.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC5, 0x30, 0x30)
        else:
            run_parent.add_run(part)


def add_code_block(doc: Document, code: str):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1A202C")
    pPr.append(shd)
    run = para.add_run(code)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0xF1, 0xF5, 0xF9)


def add_blockquote(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.35)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "3182CE")
    pBdr.append(left)
    pPr.append(pBdr)
    clean = re.sub(r"^>\s*", "", text.strip())
    run = para.add_run(clean)
    run.italic = True
    run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)


def parse_table_row(line: str):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def render_table(doc: Document, rows: list):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= col_count:
                break
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            apply_inline_formatting(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(9)
            if r_idx == 0:
                set_cell_bg(cell, "EDF2F7")
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


# ── main parser ───────────────────────────────────────────────────────────────

def parse_md(doc: Document, md_text: str):
    lines = md_text.splitlines()
    i = 0
    in_code   = False
    code_buf  = []
    table_buf = []

    while i < len(lines):
        line = lines[i]

        # fenced code block
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_buf))
                code_buf = []
                in_code  = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # flush pending table
        if table_buf and not line.strip().startswith("|"):
            render_table(doc, table_buf)
            table_buf = []

        # markdown table row
        if line.strip().startswith("|"):
            cells = parse_table_row(line)
            if all(re.match(r"^-+$", c) for c in cells if c):
                i += 1
                continue
            table_buf.append(cells)
            i += 1
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = re.sub(r"\s*---$", "", m.group(2).strip())
            if level == 1:
                para = doc.add_heading(text, level=1)
                para.paragraph_format.space_before = Pt(18)
            elif level == 2:
                para = doc.add_heading(text, level=2)
                para.paragraph_format.space_before = Pt(12)
            elif level == 3:
                para = doc.add_heading(text, level=3)
                para.paragraph_format.space_before = Pt(8)
            else:
                para = doc.add_heading(text, level=4)
            i += 1
            continue

        # horizontal rule
        if re.match(r"^---+\s*$", line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # blockquote
        if line.strip().startswith(">"):
            add_blockquote(doc, line)
            i += 1
            continue

        # table of contents entries (skip)
        if re.match(r"^- \[", line.strip()):
            i += 1
            continue

        # bullet list
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            text   = m.group(2).strip()
            style  = "List Bullet 2" if indent > 0 else "List Bullet"
            para = doc.add_paragraph(style=style)
            apply_inline_formatting(para, text)
            i += 1
            continue

        # numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            text   = m.group(2).strip()
            style  = "List Number 2" if indent > 0 else "List Number"
            para = doc.add_paragraph(style=style)
            apply_inline_formatting(para, text)
            i += 1
            continue

        # blank line
        if not line.strip():
            i += 1
            continue

        # regular paragraph
        para = doc.add_paragraph()
        apply_inline_formatting(para, line.strip())
        i += 1

    # flush if file ended mid-block
    if in_code and code_buf:
        add_code_block(doc, "\n".join(code_buf))
    if table_buf:
        render_table(doc, table_buf)


# ── document assembly ─────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("PSAT Developer Guide")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("CycleRAP v2.11  ·  Version 3  (Updated June 2026)")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    sub_para.paragraph_format.space_after = Pt(24)

    add_horizontal_rule(doc)
    doc.add_page_break()

    for path in MD_FILES:
        md_text = path.read_text(encoding="utf-8")
        parse_md(doc, md_text)
        doc.add_page_break()

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build_doc()
